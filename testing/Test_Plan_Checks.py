""" Review with Document

    Run basic plan integrity checks and parse the log file. Meant to be run
    on completed plans.

    Scope: Requires RayStation beamset to be loaded

    Example Usage:

    Script Created by RAB May 1st 2022
    Prerequisites:

    Version history:


    This program is free software: you can redistribute it and/or modify it
    under
    the terms of the GNU General Public License as published by the Free
    Software
    Foundation, either version 3 of the License, or (at your option) any later
    version.

    This program is distributed in the hope that it will be useful, but WITHOUT
    ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or
    FITNESS
    FOR A PARTICULAR PURPOSE. See the GNU General Public License for more
    details.

    You should have received a copy of the GNU General Public License along with
    this program. If not, see <http://www.gnu.org/licenses/>.
    """

__author__ = 'Adam Bayliss'
__contact__ = 'rabayliss@wisc.edu'
__date__ = '2022-Aug-03'
__version__ = '0.0.0'
__status__ = 'Testing'
__deprecated__ = False
__reviewer__ = 'Someone else'
__reviewed__ = 'YYYY-MM-DD'
__raystation__ = '11B'
__maintainer__ = 'One maintainer'
__email__ = 'rabayliss@wisc.edu'
__license__ = 'GPLv3'
__help__ = ''
__copyright__ = 'Copyright (C) 2022, University of Wisconsin Board of Regents'
__credits__ = ['']

import subprocess
import platform
import sys
import os
import io
import logging
import PySimpleGUI as sg
import re
import tkinter as tk
from docx import Document
from docx.shared import Cm, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from PIL import ImageGrab
from datetime import datetime
from collections import namedtuple, OrderedDict
from System import Environment

sys.path.insert(1, os.path.join(os.path.dirname(__file__), r'../library'))
sys.path.insert(1, os.path.join(os.path.dirname(__file__),
                                r'OldPlanReview'))
from ReviewDefinitions import *
import ExamTests
import BeamSetReviewTests
import PlanReviewTests
from GeneralOperations import find_scope

"""

 TODO: Check Beamsets for same machine

 TODO: Check for same iso, and same number of fractions in
   different beamsets, and flag for merge

 TODO:
   Check bad regions of Frame

 TODO:
   def check_plan_name(bs):
     Check plan name for appropriate
     Measure target length of prostate for pros

 TODO: Look for big gaps between targets
   def check_target_spacing(bs):
     Find all targets
     Put a box around them
     look at the gaps and if they exceed some threshold throw an alert

 TODO: If beamsets are approved
    Check the Entrance/Exit is blocked on some things
    Check that treat settings are used/appropriate
    
 TODO: Tomo Time Check
   def check_tomo_time(bs):
     Look at the plan type. Use the normal tomo mod factors
     Abdomen; 1.6 - 2.4
     Brain; 1.6 - 2.4
     Breast; 2.4 - 2.8
     Cranio - Spinal; 1.8 - 2.2
     Extremity; 2.0 - 2.4
     Gyn; 1.8 - 2.4
     H & N; 2.2 - 2.6
     Lung(non - SBRT); 2.4 - 2.8
     Lung(SBRT); 1.2 - 1.4
     Pelvis; 1.8 - 2.4
     Prostate(low; risk)    1.6 - 2.2
     Prostate(high; risk)    2.0 - 2.4


 TODO: Check collisions
   put a circle down at isocenter equal in dimension to ganty (collimator 
   pin)/bore clearance
   union patient/supports
   determine gantry positions

 TODO:
   def - check the front edges of the couch and suspended headboard

 TODO:
   Flag all ROIs not made in MIM with goals

 TODO: Stray voxel check/

 TODO: Check clinical goal
   if a clinical goal is not met, look at the objective list to see if it is 
   constrained
   
 TODO: Add test on currently commissioned beams for timestamp
 
 TODO: Check if an arc has the same couch and start/stop. if so, collimator 
 angles should differ

 TODO: FRONT PAGE CHECKS
 * TPO versus doses used in plan
 * CT Orientation
 * Number of slices and scan date
 * Special instructions
 * Energy
 
 TODO: Objective type is correct: for anything with min goals, should be 
 PTV/GTV/CTV
 
 TODO: IS there any higher order way to search for an identify a beamset
 for better tracking
 

"""


def comment_to_clipboard(rso):
    # Clear the system clipboard
    root = tk.Tk()
    root.withdraw()
    root.clipboard_clear()

    # Create the beamset comment
    approval_status = BeamSetReviewTests.approval_info(rso.plan, rso.beamset)
    beamset_comment = approval_status.beamset_approval_time

    # Copy the comment to the system clipboard
    root.clipboard_append(beamset_comment)
    root.update()

    # Return the root window so it can be destroyed by the caller
    return root


def read_log_file(patient_id):
    """
    Read the lines from the patient log in both clinical and development
    locations
    Args:
        patient_id: str: contains patient ID

    Returns:
        file_contents: lines of file
    """
    log_file = f"{patient_id}.txt"
    log_input_file = os.path.join(LOG_DIR, patient_id, log_file)
    dev_log_file = f"{patient_id}.txt"
    dev_log_input_file = os.path.join(DEV_LOG_DIR, patient_id, dev_log_file)

    file_contents = []

    try:
        with open(log_input_file) as f:
            file_contents += f.readlines()
    except FileNotFoundError:
        logging.debug(f"File {log_file} not found in dir {LOG_DIR}")

    try:
        with open(dev_log_input_file) as f:
            file_contents += f.readlines()
    except FileNotFoundError:
        logging.debug(f"File {dev_log_file} not found in dir {DEV_LOG_DIR}")

    return file_contents


def parse_log_file(lines, parent_key, phrases=KEEP_PHRASES):
    """
    Parse the log file lines for the phrase specified
    Args:
        parent_key: The top key for these log entries (typically patient level)
        lines: list of strings from a log file
        phrases: list of tuples
            (level,phrase):
                           level: is a string indicating pass level
                           phrase: is a string to identify lines for return

    Returns:
        message: list of lists of format: [parent key, key, value, result]

    Test Patient:
        Script_Testing^Plan_Review, #ZZUWQA_ScTest_01May2022, Log_Parse_Check
    """
    message = []
    time_stamp_exp = r'(^\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2})(.*)'  # The
    # front timestamp
    re_c = r'^\t(Case: .*)\t'
    re_e = r'(Exam: .*)\t'
    re_p = r'(Plan: .*)\t'
    re_b = r'(Beamset: [^\t|\s]+)'  # Terminate after this level

    # Declare reg-ex for levels in the log file
    context_exps = OrderedDict()
    context_exps['Beamset'] = re_c + re_e + re_p + re_b
    context_exps['Plan'] = re_c + re_e + re_p
    context_exps['Exam'] = re_c + re_e
    context_exps['Case'] = re_c

    for p in phrases:
        key = 'Log' + p[1]
        message.append([parent_key, key, key, p[0]])
        re_phrase = p[1] + r'.*\.py: '

        for line in lines:
            if p[1] in line:
                # Remove the source python program from the line
                line = re.sub(re_phrase, '', line)

                # Sort the line into a part for the timestamp and one for
                # remainder
                parsed_l = [part for t in re.findall(time_stamp_exp, line) for
                            part in t]
                parsed_l[1] = parsed_l[
                    1].lstrip().rstrip()  # Remove front and back white space

                deepest_level = None
                for c, exp in context_exps.items():
                    if bool(re.search(exp, parsed_l[1])):
                        levels = OrderedDict()
                        deepest_level = c
                        for g in re.findall(exp, parsed_l[1])[0]:
                            lev_key, lev_val = re.split(': ', g, maxsplit=1)
                            levels[lev_key] = lev_val
                        parsed_l[1] = re.sub(exp, '', parsed_l[1])
                        parsed_l[0] += ' ' + deepest_level + ': ' + levels[
                            deepest_level]
                        break

                if not deepest_level:
                    parsed_l[1] = re.sub(r'\t', '', parsed_l[1])

                message.append([key, parsed_l[0], parsed_l[0], parsed_l[1]])

    return message


def parse_treatment_planning_order_selection(beamset_name, messages,
                                             dialog_key):
    """
    Using the beamset name and the log file, use regex to find a phrase
    identifying the
    template name used in the treatment planning order for this beamset.
    Args:
        beamset_name: <str> name of current beamset
        messages: [ <str> ]: list of strings
        dialog_key: key to be used for top level entry in the tree

    Returns:

    TODO: Take a reg-exp as a list for input for matching a dialog and for
    each desired phrase
          loop over the phrases for a match
    TODO: Add the target matching that takes place for this step with
    consideration of the
          pre-logcrit syntax and post-logcrit syntax


    """
    # Parse the dialogs for specific key phrases related to the beamset dialog
    beamset_template_searches = {'Dialog': re.compile(
        r'(Treatment Planning Order selected:\s?)(.*)$')}
    template_data = {dialog_key: (ALERT,
                                  f'Treatment Planning for Beamset '
                                  f'{beamset_name} goals manually defined')}
    for m in messages:
        template_search = re.search(beamset_template_searches['Dialog'], m[3])
        if template_search and beamset_name in m[1]:
            # Found the TPO Dialog. Lets display it
            # Note that it is a little sloppy, since this will always grab
            # the last match in the log file
            k, template_name = template_search.groups()
            template_data[dialog_key] = (None, template_name)
    return template_data


def parse_beamset_selection(beamset_name, messages, dialog_key):
    # Go in here and if the message line matches the beamset_name
    # then return all the choices in this dialog as a list of tuples

    # Parse the dialogs for specific key phrases related to the beamset dialog
    beamset_template_searches = {
        'Dialog': re.compile(r'(Dialog):\s?(Beamset Template Selection)'),
        'Name': re.compile(r'(TemplateName):\s?(.*)Iso'),
        'NameAlt': re.compile(r'(TemplateName):\s?([^\t]+)'),
        'Isocenter': re.compile(r'(Iso):\s?([^\t]+)'),
        'Energy': re.compile(r'(Energy):\s?([^\t]+)')}
    # Find all the message lines that relate to a beamset dialog
    dialog_key = 'Beamset Template Selection'
    template_data = {dialog_key: (
        ALERT, 'Beamset {} not set by script'.format(beamset_name))}
    # loop over all lines in the template
    for m in messages:
        if re.search(beamset_template_searches['Dialog'],
                     m[3]) and beamset_name in m[1]:
            # This is a beamset dialog: add it to the beamset listings
            # Get the template name
            try:
                k, template_name = re.search(beamset_template_searches['Name'],
                                             m[3]).groups()
            except AttributeError:
                k, template_name = re.search(
                    beamset_template_searches['NameAlt'], m[3]).groups()
            # Get the isocenter information
            k, iso_info = re.search(beamset_template_searches['Isocenter'],
                                    m[3]).groups()
            k, energy = re.search(beamset_template_searches['Energy'],
                                  m[3]).groups()
            template_data[dialog_key] = (None, m[1])
            template_data['Template Name'] = (None, template_name)
            # TODO: Replace with a parse on bracketed entries return
            #  placement method
            template_data['Isocenter'] = (None, iso_info)
            template_data['Energy'] = (None, energy)
    return template_data


def NEW_parse_beamset_selection(beamset_name, messages):
    """
    Parse the messages for a specific beamset dialog and return the dialog
    choices as a dictionary.

    Args:
        beamset_name (str): The name of the beamset to look for in the messages.
        messages (list): A list of messages to search through.

    Returns:
        A dictionary containing the beamset template dialog choices.
    """

    # Define regular expressions to search for in the messages
    beamset_template_searches = {
        'Dialog': re.compile(r'(Dialog):\s?(Beamset Template Selection)'),
        'Name': re.compile(r'(TemplateName):\s?(.*)Iso'),
        'NameAlt': re.compile(r'(TemplateName):\s?([^\t]+)'),
        'Isocenter': re.compile(r'(Iso):\s?([^\t]+)'),
        'Energy': re.compile(r'(Energy):\s?([^\t]+)')
    }

    # Initialize the template_data dictionary with default values
    template_data = {
        'Beamset Template Selection': (
            None, 'Beamset {} not set by script'.format(beamset_name)),
        'Template Name': (None, ''),
        'Isocenter': (None, ''),
        'Energy': (None, '')
    }

    # Search through the messages for a beamset dialog matching the beamset name
    for message in messages:
        if re.search(beamset_template_searches['Dialog'],
                     message[3]) and beamset_name in message[1]:
            # Found a matching beamset dialog: extract the template information
            try:
                _, template_name = re.search(beamset_template_searches['Name'],
                                             message[3]).groups()
            except AttributeError:
                _, template_name = re.search(
                    beamset_template_searches['NameAlt'], message[3]).groups()
            _, iso_info = re.search(beamset_template_searches['Isocenter'],
                                    message[3]).groups()
            _, energy = re.search(beamset_template_searches['Energy'],
                                  message[3]).groups()

            # Update the template_data dictionary with the extracted information
            template_data['Beamset Template Selection'] = (None, message[1])
            template_data['Template Name'] = (None, template_name)
            # TODO: Replace with a parse on bracketed entries return
            #  placement method
            template_data['Isocenter'] = (None, iso_info)
            template_data['Energy'] = (None, energy)

    return template_data


def build_tree_element(parent_key: str, child_key: str, pass_result: str,
                       message_str: str) -> tuple:
    """
    Builds an element for insertion in a PySimpleGui tree.

    Args:
        parent_key (str): The key of the parent node.
        child_key (str): The key of the child node.
        pass_result (str): The result of the test, either "Pass", "Fail",
        "Alert", or "".
        message_str (str): The user message.

    Returns:
        A tuple containing two lists.
        Parent list contains the node for the node and child
        list is the sub_node. Each list contains the following elements:
        [key, value, text, pass/fail result, icon]
    """

    # Determine the appropriate icon based on the pass result
    if pass_result == FAIL:
        icon = RED_CIRCLE
    elif pass_result == PASS:
        icon = GREEN_CIRCLE
    elif pass_result == ALERT:
        icon = YELLOW_CIRCLE
    else:
        icon = BLUE_CIRCLE

    # Create the node for the parent
    parent_node = [parent_key, child_key, child_key, pass_result, icon]

    # Create the node for the child
    child = [child_key, parent_key + '.' + child_key, message_str,
             pass_result, icon]

    return parent_node, child


def parse_level_tests(level_tests):
    # Insert Level Nodes
    overall_status = PASS
    status_icon = GREEN_CIRCLE

    statuses = [level[3] for level in level_tests]
    if FAIL in statuses:
        overall_status = FAIL
        status_icon = RED_CIRCLE
    elif ALERT in statuses:
        overall_status = ALERT
        status_icon = YELLOW_CIRCLE

    return overall_status, status_icon


def insert_tests_return_fails(test_data, treedata, fails):
    """
    Inserts test data into PySimpleGUI.TreeData and returns a list of test
    failures.

    Args:
        test_data (list): A list of test data to be inserted into the tree.
        Each element should be
            a tuple in the format (test_name, test_key, test_description,
            test_result, test_icon).
        treedata (PySimpleGUI.TreeData): The PySimpleGUI TreeData object to
        insert the test data into.
        fails (list): An empty list that will be populated with any test
        failures.

    Returns:
        bool: True if test_data is not empty, False otherwise.
    """
    if test_data:
        for test in test_data:
            test_name, test_key, test_description, test_result, test_icon = test
            treedata.Insert(test_name, test_key, test_description,
                            [test_result, ""], icon=test_icon)
            if test_key == FAIL or test_key == ALERT:
                fails.append(
                    f"TEST: {test_name}: {test_key}: {test_description}")
        return True
    else:
        return False


def get_exam_level_tests(rso):
    #
    # Get target length
    target_extent = ExamTests.get_si_extent(rso, types=['Ptv'])
    patient_checks_dict = {
        "DICOM RayStation Comparison":
            (ExamTests.check_exam_data, {}),
        "Exam Date Is Recent":
            (ExamTests.compare_exam_date, {}),
        "Localization Point Exists":
            (ExamTests.check_localization, {}),
        "Contours are interpolated":
            (ExamTests.check_contour_gaps, {}),
        "Supports correctly overriden":
            (ExamTests.check_support_material, {}),
        "Image Is Axially Oriented":
            (ExamTests.match_image_directions, {}),
    }
    # TODO: If the target extent is NONE, then we ought to try and get one
    #  from dose
    if target_extent:
        patient_checks_dict.update({
            "Image extent sufficient":
                (ExamTests.image_extent_sufficient,
                 {'TARGET_EXTENT': target_extent}),
            "Couch extent sufficient":
                (ExamTests.couch_extent_sufficient,
                 {'TARGET_EXTENT': target_extent}),
            "Edge of scan overlaps patient at key slices":
                (ExamTests.external_overlaps_fov,
                 {'TARGET_EXTENT': target_extent}),
        })
    return patient_checks_dict


def get_plan_level_tests(rso, physics_review=True):
    plan_checks_dict = {
        "Plan approval status":
            (PlanReviewTests.check_plan_approved,
             {"do_physics_review": physics_review})
    }
    return plan_checks_dict


def get_beamset_level_tests(rso, physics_review=True):
    beamset_checks_dict = {
        "Beamset approval status":
            (BeamSetReviewTests.check_beamset_approved,
             {"do_physics_review": physics_review}),
        "Isocenter Position Identical":
            (BeamSetReviewTests.check_common_isocenter, {"tolerance": 1e-15}),
        "Check Fractionation":
            (BeamSetReviewTests.check_fraction_size, {}),
        "Slice Thickness Comparison":
            (BeamSetReviewTests.check_slice_thickness, {}),
        "Bolus Application":
            (BeamSetReviewTests.check_bolus_included, {}),
        "No Fly Zone Dose Check":
            (BeamSetReviewTests.check_no_fly, {}),
        "Check for pacemaker compliance":
            (BeamSetReviewTests.check_pacemaker, {}),
        "Dose Grid Size Check":
            (BeamSetReviewTests.check_dose_grid, {}),
        "Planning Risk Volume Assessment":
            (BeamSetReviewTests.check_prv_status, {}),
        "Couch Zero Full Rotation Clearance Check":
            (BeamSetReviewTests.check_isocenter_clearance, {}),
    }

    # Plan check for VMAT
    #
    technique = rso.beamset.DeliveryTechnique if rso.beamset else None
    if technique == 'DynamicArc':
        if rso.beamset.Beams[0].HasValidSegments:
            beamset_checks_dict["Control Point Spacing"] = (
                BeamSetReviewTests.check_control_point_spacing,
                {'expected': 2.})
            beamset_checks_dict["Beamset Complexity"] = (
                BeamSetReviewTests.compute_beam_properties, {})
    elif technique == 'SMLC':
        try:
            _ = rso.beamset.Beams[0].Segments[
                0]  # Determine if beams have segments
            beamset_checks_dict["Beamset Complexity"] = (
                BeamSetReviewTests.compute_beam_properties, {})
            beamset_checks_dict["EDW MU Check"] = (
                BeamSetReviewTests.check_edw_MU, {})
            beamset_checks_dict["EDW FieldSize Check"] = (
                BeamSetReviewTests.check_edw_field_size, {})
        except Exception as e:
            logging.debug('Cannot check beamsets yet {}'.format(str(e)))
    elif 'Tomo' in technique:
        try:
            _ = rso.beamset.Beams[0].Segments[
                0]  # Determine if beams have segments
            beamset_checks_dict["Isocenter Lateral Acceptable"] = (
                BeamSetReviewTests.check_tomo_isocenter, {})
            beamset_checks_dict["Modulation Factor Acceptable"] = (
                BeamSetReviewTests.check_mod_factor, {})
            beamset_checks_dict["Transfer BeamSet Approval Status"] = (
                BeamSetReviewTests.check_transfer_approved, {})
        except Exception as e:
            logging.debug('Cannot check beamsets yet {}'.format(str(e)))
    return beamset_checks_dict


# EVENTS
# Create the event handler for the radio buttons
def handle_radio_event(window, event):
    if event.endswith('_RADIO_NO'):
        input_key = event.replace('_RADIO_NO', '_INPUT')
        window[input_key].update(text_color='#000000',
                                 background_color='#ffffff')
    if event.endswith('_RADIO_YNA'):
        input_key = event.replace('_RADIO_YNA', '_INPUT')
        window[input_key].update(text_color='#ffffff',
                                 background_color='#af0404')


def create_tab_manual_checks(failed_tests):
    """
    Create a PySimpleGUI tab with multiple frames, each containing a checklist of options.

    Parameters:
    -----------
    failed_tests : list
        List of qa_tests that failed for user to review. Format is [comment, result of test,
        stoplight icon]

    Returns:
    --------
    PySimpleGUI.Tab
        A PySimpleGUI Tab object.
    """

    # Determine the maximum length of any checklist item
    max_check = max([len(item['text']) for key in CHECK_BOXES for item in CHECK_BOXES[key]])

    # Create the list of elements for each key in CHECK_BOXES
    layout = [
        [sg.Text('Select an option for each item:')]
    ]
    for key in CHECK_BOXES:
        # Create a frame for the current key
        frame_layout = [
            [
                # Checklist item text
                sg.Text(item['text'], size=(max_check, 2)),

                # Radio buttons for Yes/NA and No options
                sg.Radio('Yes/NA', f'{item["key"]}_RADIO',
                         default=False, key=f'{item["key"]}_RADIO_YNA',
                         enable_events=True),
                sg.Radio('No', f'{item["key"]}_RADIO',
                         default=False, key=f'{item["key"]}_RADIO_NO',
                         enable_events=True),

                # Input text field for additional comments
                sg.InputText(size=(int(max_check * 1.2), 1),
                             key=f'{item["key"]}_INPUT',
                             enable_events=True)
            ]
            for item in CHECK_BOXES[key]
        ]
        frame = sg.Frame(key, frame_layout)
        layout.append([frame])

    # Create frames for failed qa_tests
    rows = []
    for comment, result, icon in failed_tests:
        rows.append([sg.Image(icon),
                     sg.Text(result,
                             size=(max_check, None),
                             auto_size_text=True,
                             justification='left'),
                     sg.InputText(
                         default_text=f"{comment}: Comment",
                         key=comment,
                         size=(1.2 * max_check, None),
                         enable_events=True,
                         pad=((40, 0), (0, 0)),
                         text_color='#000000',
                         background_color='#ffffff',
                         border_width=0,
                         justification='left',
                         tooltip=comment)
                     ])

    frame_failed_tests = sg.Frame('Failed Tests', [*rows])
    layout.append([frame_failed_tests])

    # Create and return the tab
    tab = sg.Tab('Miller Time', layout)
    return tab


def capture_screen(window):
    # Determine the operating system
    os_name = platform.system()

    # Take a screenshot based on the operating system
    if os_name == 'Windows':
        # Run the Windows Snipping Tool and wait for it to finish
        subprocess.run(["SnippingTool.exe"])
        # Get the image data from the clipboard
        img = ImageGrab.grabclipboard()

        # Convert the image to bytes
        with io.BytesIO() as img_bytes:
            img.save(img_bytes, format='PNG')
            img_data = img_bytes.getvalue()

    else:
        # Unsupported operating system
        raise NotImplementedError(
            f'Operating system "{os_name}" is not supported.')

    return img_data


def save_report(patient_id, beamset_name, user_name, description, screenshot):
    """
    Save the error report document.

    Parameters:
        patient_id (str): The patient ID.
        beamset_name (str): The beamset name.
        user_name (str): The user name.
        description (str): The description of the error.
            screenshot (bytes): The screenshot image data.

    Returns:
        None
    """
    # Create a filename for the report
    now = datetime.now()
    filename = f"{patient_id}_{beamset_name}_" \
               f"{now.strftime('%Y-%m-%d_%H-%M-%S')}.docx"

    # Embed the screenshot into the doc file
    doc = Document()

    # Set page orientation to landscape
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    # Set narrow margins
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    # Add report content
    doc.add_heading('Error Report')
    doc.add_paragraph(f'Patient ID: {patient_id}')
    doc.add_paragraph(f'Beamset Name: {beamset_name}')
    doc.add_paragraph(f'User Name: {user_name}')
    doc.add_paragraph(f'Time:{now.strftime("%Y-%m-%d_%H-%M-%S")}')
    doc.add_paragraph(f'Description of Error: {description}')
    doc.add_picture(io.BytesIO(screenshot), width=Cm(25))

    # Save the document
    doc.save(os.path.join(ERROR_DIR, filename))

    # Show success message
    sg.popup(f'Error report saved as {filename}')


def report_error(rso):
    # Define the layout of the error report dialog
    try:
        user_name = str(Environment.UserName)
    except Exception as e:
        logging.debug('{}'.format(e))
        user_name = None

    error_report_layout = [
        [sg.Text('Patient ID'), sg.Input(default_text=rso.patient.PatientID,
                                         key='patient_id')],
        [sg.Text('Beamset Name:'),
         sg.Input(default_text=rso.beamset.DicomPlanLabel,
                  key='beamset_name')],
        [sg.Text('User Name:'),
         sg.Input(default_text=user_name, key='user_name')],
        [sg.Text('Description:')],
        [sg.Multiline(key='description', size=(50, 10))],
        [sg.Button("Capture",
                   tooltip='Capture a screenshot with Snipping Tool: select '
                           '"New",'
                           + ' capture your screen, and press "Ctrl-C" to '
                             'save to clipboard.'),
         sg.Button("Finish")],
    ]

    # Create the dialog window
    error_report_window = sg.Window('Error Report', error_report_layout)
    img_data = None
    # Event loop for the dialog window
    while True:
        event, values = error_report_window.read()
        if event == sg.WIN_CLOSED:
            break
        elif event == 'Capture':
            # Take a screenshot
            img_data = capture_screen(error_report_window)
            if img_data:
                sg.popup_ok('Screenshot captured!')
            else:
                sg.popup_ok(
                    'Oops I missed it. Try hitting Ctrl-C after you capture')
                img_data = capture_screen(error_report_window)
        elif event == 'Finish':
            # Save the report and close the window
            patient_id = values['patient_id']
            beamset_name = values["beamset_name"]
            user_name = values["user_name"]
            description = values["description"]
            screenshot = img_data
            save_report(patient_id, beamset_name, user_name, description,
                        screenshot)
            error_report_window.close()
            break


def check_plan(physics_review=True, rso=None):
    """
        patient_key
            |
             -- Patient Checks
            |
             -- exam_key
                    |
                     -- DICOM Checks
            |
             -- structure_key
                    |
                     -- Structure Checks
             -- plan_key
                    |
                     -- Plan Checks
             - beamset_key
                    |
                     -- Beamset Checks
            |
             -- Logs
    """

    #
    try:
        user_name = str(Environment.UserName)
    except Exception as e:
        logging.debug('{}'.format(e))
        user_name = None

    tree_data = sg.TreeData()

    if not rso:
        # Initialize return variable
        Pd = namedtuple('Pd', ['error', 'db', 'case', 'patient', 'exam', 'plan',
                               'beamset'])
        # Get current patient, case, exam
        rso = Pd(error=[],
                 patient=find_scope(level='Patient'),
                 case=find_scope(level='Case'),
                 exam=find_scope(level='Examination'),
                 db=find_scope(level='PatientDB'),
                 plan=find_scope(level='Plan'),
                 beamset=find_scope(level='BeamSet'))
    r = comment_to_clipboard(rso)
    #
    # Tree Levels
    patient_key = (LEVELS['PATIENT_KEY'], "Patient: " + rso.patient.PatientID)
    exam_key = (LEVELS['EXAM_KEY'], "Exam: " + rso.exam.Name)
    plan_key = (LEVELS['PLAN_KEY'], "Plan: " + rso.plan.Name)
    beamset_key = (
        LEVELS['BEAMSET_KEY'], "Beam Set: " + rso.beamset.DicomPlanLabel)
    rx_key = (LEVELS['RX_KEY'], "Prescription")
    log_key = (LEVELS['LOG_KEY'], "Logging")
    #

    test_results = []
    tree_children = []

    """
    Gather Patient Level Checks
    """
    if rso.exam:
        patient_checks_dict = get_exam_level_tests(rso)
    check_patient_logs = True if rso.patient else False
    plan_checks_dict = {}
    """
    Gather Plan Level Checks
    """
    if rso.plan:
        plan_checks_dict = get_plan_level_tests(rso, physics_review)
    """
    Gather BeamSet Level Checks
    """
    if rso.beamset:
        beamset_checks_dict = get_beamset_level_tests(rso, physics_review)
    """
    Parse logs
    """
    if check_patient_logs:
        lines = read_log_file(patient_id=rso.patient.PatientID)
        message_logs = parse_log_file(lines=lines, parent_key=log_key[0],
                                      phrases=KEEP_PHRASES)
    # Execute qa_tests
    exam_level_tests = []
    for key, p_func in patient_checks_dict.items():
        pass_result, message = p_func[0](rso=rso, **p_func[1])
        node, child = build_tree_element(parent_key=exam_key[0],
                                         child_key=key,
                                         pass_result=pass_result,
                                         message_str=message)
        exam_level_tests.extend([node, child])
        tree_children.append(child)

    """
    Execute Plan Level Checks
    """
    # Plan LevelChecks
    plan_level_tests = []
    for key, pl_func in plan_checks_dict.items():
        pass_result, message = pl_func[0](rso=rso, **pl_func[1])
        node, child = build_tree_element(parent_key=plan_key[0],
                                         child_key=key,
                                         pass_result=pass_result,
                                         message_str=message)
        plan_level_tests.extend([node, child])
        tree_children.append(child)

    #
    # BEAMSET LEVEL CHECKS
    beamset_level_tests = []

    #
    # Run dialog parse
    dialog_key = 'Beamset Template Selection'
    beamset_dialog = parse_beamset_selection(
        beamset_name=rso.beamset.DicomPlanLabel,
        messages=message_logs,
        dialog_key=dialog_key)
    node, child = build_tree_element(parent_key=beamset_key[0],
                                     child_key=dialog_key,
                                     pass_result=beamset_dialog[dialog_key][0],
                                     message_str=beamset_dialog[dialog_key][1])
    beamset_level_tests.extend([node, child])
    for k, v in beamset_dialog.items():
        if k != dialog_key:
            node, child = build_tree_element(parent_key=dialog_key,
                                             child_key=k,
                                             pass_result=v[0],
                                             message_str=v[1])
            beamset_level_tests.extend([node, child])
            tree_children.append(child)
    # Read the treatment planning order selection
    dialog_key = 'Treatment Planning Order Selection'
    tpo_dialog = parse_treatment_planning_order_selection(
        beamset_name=rso.beamset.DicomPlanLabel,
        messages=message_logs,
        dialog_key=dialog_key)
    node, child = build_tree_element(parent_key=beamset_key[0],
                                     child_key=dialog_key,
                                     pass_result=tpo_dialog[dialog_key][0],
                                     message_str=tpo_dialog[dialog_key][1])
    beamset_level_tests.extend([node, child])
    for k, v in tpo_dialog.items():
        if k != dialog_key:
            node, child = build_tree_element(parent_key=dialog_key,
                                             child_key=k,
                                             pass_result=v[0],
                                             message_str=v[1])
            beamset_level_tests.extend([node, child])
            tree_children.append(child)
    # Run others
    for key, b_func in beamset_checks_dict.items():
        pass_result, message = b_func[0](rso=rso, **b_func[1])
        node, child = build_tree_element(parent_key=beamset_key[0],
                                         child_key=key,
                                         pass_result=pass_result,
                                         message_str=message)
        beamset_level_tests.extend([node, child])
        tree_children.append(child)

    # TREE BUILDING
    #
    # Insert main (patient) node
    patient_level_pass, patient_icon = parse_level_tests(
        exam_level_tests + plan_level_tests + beamset_level_tests)
    tree_data.Insert("", patient_key[0], patient_key[1], patient_level_pass,
                     icon=patient_icon)
    # Insert Exam Level Nodes
    exam_level_pass, exam_icon = parse_level_tests(exam_level_tests)
    tree_data.Insert(patient_key[0], exam_key[0], exam_key[1], exam_level_pass,
                     icon=exam_icon)
    insert_tests_return_fails(exam_level_tests, tree_data, test_results)
    # Insert Plan Level notes
    plan_level_pass, plan_icon = parse_level_tests(plan_level_tests)
    tree_data.Insert(patient_key[0], plan_key[0], plan_key[1], plan_level_pass,
                     icon=plan_icon)
    insert_tests_return_fails(plan_level_tests, tree_data, test_results)
    #
    # Insert Beamset Level Nodes
    beamset_level_pass, beamset_icon = parse_level_tests(beamset_level_tests)
    tree_data.Insert(patient_key[0], beamset_key[0], beamset_key[1],
                     beamset_level_pass, icon=beamset_icon)
    insert_tests_return_fails(beamset_level_tests, tree_data, test_results)
    #
    # Log Level
    tree_data.Insert(patient_key[0], log_key[0], log_key[1], "")
    if check_patient_logs:
        for m in message_logs:
            tree_data.Insert(m[0], m[1], m[2], [
                m[3]])  # Note the list of the last entry. Can this be of use?

    # Gui
    # TODO is it possible to specify which columns are expanded by pass fail?
    sg.theme('DarkBrown4')
    col_widths = [20]
    comment_visible = False
    bottom = [[sg.Button('Cancel'),
               sg.Button('Done'),
               sg.Button('Report Error', key='report_error'),
               sg.Frame('',
                        [[sg.InputText('Report script error here',
                                       key='-COMMENT TEXT-'),
                          sg.Button('Ok')]],
                        visible=comment_visible,
                        key='-COMMENT TEXT VISIBLE-'
                        ),
               ]]
    tab1 = [[sg.Frame('ReviewChecks:',
                      [
                          [
                              sg.Tree(
                                  data=tree_data,
                                  headings=['Result'],
                                  auto_size_columns=False,
                                  num_rows=40,
                                  col0_width=120,
                                  col_widths=col_widths,
                                  key='-TREE-',
                                  show_expanded=False,
                                  justification="left",
                                  enable_events=True),
                          ],
                          # [Sg.Frame('', bottom)],
                      ],
                      pad=(0, 0))
             ]]
    c0 = []
    c1 = []
    # Keys are generated for the Manual Checks tab in the form:
    # CHECK_BOXES['Plan Settings']['plan_name'] + +'_RADIO_YNA'
    # Update keys for all elements of dialog
    # Failed test sorting
    failed_tests = []
    for comment, _, result, pass_fail, icon in tree_children:
        # for test in alt_test_readout:
        if pass_fail != PASS:
            c0.append([sg.Image(icon), sg.Text(result), ])
            c1.append([sg.InputText(default_text=comment + ": Comment",
                                    key=comment)])
            failed_tests.append([comment, result, icon])
    tab2 = [[sg.Column(c0), sg.Column(c1)]]
    tab3 = create_tab_manual_checks(failed_tests)

    layout = [[sg.TabGroup([[sg.Tab('Review and Logs', tab1)],
                            # [Sg.Tab('Comments', tab2)],
                            [tab3]])],
              [sg.Frame('', bottom)],
              ]

    window = sg.Window('Plan Review: ' + user_name, layout)

    while True:  # Event Loop
        event, values = window.read()
        if event in (sg.WIN_CLOSED, 'Cancel'):
            break
        elif event in 'Done':
            if test_results and not physics_review:
                now = datetime.now()
                dt_string = now.strftime("(%H:%M) %B %d, %Y")
                logging.warning(
                    "Review Script Warnings/Errors present at {}".format(
                        dt_string))
                for tr in test_results:
                    logging.warning(f"\t{tr}")
            break
        elif event.endswith('_RADIO_NO'):
            handle_radio_event(window, event)
        elif event.endswith('_RADIO_YNA'):
            handle_radio_event(window, event)
        elif event == 'report_error':
            report_error(rso)

    window.close()
    r.destroy()
    return {'Test_Exam': exam_level_tests,
            'Test_Plan': plan_level_tests,
            'Test_BeamSet': beamset_level_tests}


def generate_doc(rso, tests):
    # Output file
    file_name = rso.patient.PatientID + "_" + rso.beamset.DicomPlanLabel + \
                ".doc"
    output_file = os.path.join(OUTPUT_DIR, file_name)
    header_text = "Photon VMAT Physics Review"
    # Get approval info:
    approval_status = BeamSetReviewTests.approval_info(rso.plan, rso.beamset)
    if approval_status.beamset_approved:
        current_time = str(rso.beamset.Review.ReviewTime)
    else:
        current_time = 'NA'

    demographics = {
        'Name': rso.patient.Name,
        'MRN': rso.patient.PatientID,
        'Beamset Name': rso.beamset.DicomPlanLabel,
        'Approval Time': current_time}
    # Responses to non-scriptable questions
    plan_questions = tests['Test_BeamSet']
    # plan_questions = [('Plan Name consistent with TPO', 'Y', 'Sample Input'),
    #                   ('Plan approved by MD', 'NA', 'Preplan'),
    #                   ]

    # Begin
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    # Add header

    header = section.header
    paragraph = header.paragraphs[0]
    # Add logo
    logo_run = paragraph.add_run()
    logo_run.add_picture(UW_HEALTH_LOGO, width=Inches(1.0))
    text_run = paragraph.add_run()
    text_run.text = '\t' + header_text  # For center align of text
    text_run.style = "Heading 1 Char"

    paragraph = document.add_paragraph()
    # Add Top Row Demographics
    table = document.add_table(rows=2, cols=4, style='Medium Grid 1 Accent 2')
    for index, k in enumerate(demographics):
        row_key = table.rows[0]
        row_value = table.rows[1]
        row_key.cells[index].text = k
        row_value.cells[index].text = demographics[k]
    # Add the plan checks
    paragraph = document.add_paragraph()
    document = add_check_list_table(tests['Test_BeamSet'], document)

    document.save(output_file)
    print('Complete')


def add_check_list_table(check_results, document, title=None):
    logging.debug(f'{check_results}')
    n_cols = 3  # Icon, Testname, Result, Comment
    table_properties = {'NCOL': 4,
                        'WIDTH_COL': [(0, 0.25), (1, 1.), (2, 3.), (3, 3.)]}
    table = document.add_table(rows=1, cols=table_properties['NCOL'],
                               style='Light Grid Accent 2')
    i = 0
    for r in enumerate(check_results):
        row = table.rows[i]
        child_list = r[1]
        if i == 0:
            row.cells[0].text = 'Status'
            row.cells[1].text = 'Test Performed'
            row.cells[2].text = 'Result'
            row.cells[3].text = 'Reviewer Comment'
            table.add_row()
            i += 1
        elif child_list[0] not in LEVELS.values():
            logging.debug('Child list {}'.format(child_list))
            row.cells[0].add_paragraph().add_run().add_picture(child_list[4],
                                                               width=Inches(
                                                                   0.2),
                                                               height=Inches(
                                                                   0.2))
            row.cells[1].text = child_list[0]
            row.cells[2].text = child_list[2]
            table.add_row()
            i += 1
        for index, width in table_properties['WIDTH_COL']:
            for cell in table.columns[index].cells:
                cell.width = Inches(width)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return document


def main(physics_review=True):
    # Initialize return variable
    Pd = namedtuple('Pd', ['error', 'db', 'case', 'patient', 'exam', 'plan',
                           'beamset'])
    # Get current patient, case, exam
    rso = Pd(error=[],
             patient=find_scope(level='Patient'),
             case=find_scope(level='Case'),
             exam=find_scope(level='Examination'),
             db=find_scope(level='PatientDB'),
             plan=find_scope(level='Plan'),
             beamset=find_scope(level='BeamSet'))
    tests = check_plan(physics_review, rso=rso)
    if physics_review:
        generate_doc(rso, tests=tests)


if __name__ == '__main__':
    main(physics_review=False)
