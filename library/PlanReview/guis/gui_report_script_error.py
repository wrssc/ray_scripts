import subprocess
import platform
from PIL import ImageGrab
import io
import os
import PySimpleGUI as sg
from datetime import datetime
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm
from PlanReview.utils import get_user_name
from PlanReview.review_definitions import ERROR_DIR


def capture_screen(window):
    # Determine the operating system
    os_name = platform.system()

    # Take a screenshot based on the operating system
    if os_name == 'Windows':
        # Run the Windows Snipping Tool and wait for it to finish
        subprocess.run(["SnippingTool.exe"])
        # Get the image data from the clipboard
        img = ImageGrab.grabclipboard()

        # Convert the image to bytes
        with io.BytesIO() as img_bytes:
            img.save(img_bytes, format='PNG')
            img_data = img_bytes.getvalue()

    else:
        # Unsupported operating system
        raise NotImplementedError(
            f'Operating system "{os_name}" is not supported.')

    return img_data


def save_report(patient_id, beamset_name, user_name, description, screenshot):
    """
    Save the error report document.

    Parameters:
        patient_id (str): The patient ID.
        beamset_name (str): The beamset name.
        user_name (str): The user name.
        description (str): The description of the error.
            screenshot (bytes): The screenshot image data.

    Returns:
        None
    """
    # Create a filename for the report
    now = datetime.now()
    filename = f"{patient_id}_{beamset_name}_" \
               f"{now.strftime('%Y-%m-%d_%H-%M-%S')}.docx"

    # Embed the screenshot into the doc file
    doc = Document()

    # Set page orientation to landscape
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    # Set narrow margins
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    # Add report content
    doc.add_heading('Error Report')
    doc.add_paragraph(f'Patient ID: {patient_id}')
    doc.add_paragraph(f'Beamset Name: {beamset_name}')
    doc.add_paragraph(f'User Name: {user_name}')
    doc.add_paragraph(f'Time:{now.strftime("%Y-%m-%d_%H-%M-%S")}')
    doc.add_paragraph(f'Description of Error: {description}')
    doc.add_picture(io.BytesIO(screenshot), width=Cm(25))

    # Save the document
    doc.save(os.path.join(ERROR_DIR, filename))

    # Show success message
    sg.popup(f'Error report saved as {filename}')


def report_script_error(rso):
    # Define the layout of the error report dialog
    user_name = get_user_name()

    error_report_layout = [
        [sg.Text('Patient ID'), sg.Input(default_text=rso.patient.PatientID,
                                         key='patient_id')],
        [sg.Text('Beamset Name:'),
         sg.Input(default_text=rso.beamset.DicomPlanLabel,
                  key='beamset_name')],
        [sg.Text('User Name:'),
         sg.Input(default_text=user_name, key='user_name')],
        [sg.Text('Description:')],
        [sg.Multiline(key='description', size=(50, 10))],
        [sg.Button("Capture",
                   tooltip='Capture a screenshot with Snipping Tool: select '
                           '"New",'
                           + ' capture your screen, and press "Ctrl-C" to '
                             'save to clipboard.'),
         sg.Button("Finish")],
    ]

    # Create the dialog window
    error_report_window = sg.Window('Error Report', error_report_layout)
    img_data = None
    # Event loop for the dialog window
    while True:
        event, values = error_report_window.read()
        if event == sg.WIN_CLOSED:
            break
        elif event == 'Capture':
            # Take a screenshot
            img_data = capture_screen(error_report_window)
            if img_data:
                sg.popup_ok('Screenshot captured!')
            else:
                sg.popup_ok(
                    'Oops I missed it. Try hitting Ctrl-C after you capture')
                img_data = capture_screen(error_report_window)
        elif event == 'Finish':
            # Save the report and close the window
            patient_id = values['patient_id']
            beamset_name = values["beamset_name"]
            user_name = values["user_name"]
            description = values["description"]
            screenshot = img_data
            save_report(patient_id, beamset_name, user_name, description,
                        screenshot)
            error_report_window.close()
            break
