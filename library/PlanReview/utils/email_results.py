import subprocess
import platform
from PIL import ImageGrab
import io
import os
import smtplib
import xml.etree.ElementTree as ET
import PySimpleGUI as Sg
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from datetime import datetime
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm
from PlanReview.review_definitions import ERROR_DIR, QI_REPORTS_DIR
from PlanReview.utils import get_user_name


def list_format(input_str):
    if isinstance(input_str, str):
        # Split the string by commas and strip whitespace
        return [addr.strip() for addr in input_str.split(',') if addr.strip()]
    elif input_str is None:
        return []
    elif isinstance(input_str, list):
        return input_str
    else:
        return []


def read_email_preferences(xml_file_path, email_type):
    """
    Read the email preferences from the XML file based on the email type.

    Parameters:
        xml_file_path (str): Path to the XML configuration file.
        email_type (str): The type of email configuration to read.

    Returns:
        dict: A dictionary containing the email configuration.
    """
    tree = ET.parse(xml_file_path)
    root = tree.getroot()

    # If there's only one <email_report>, use it
    if root.tag == 'email_report':
        report_type = root.find('report_type').text.strip()
        if report_type != email_type:
            raise ValueError(f'Email type "{email_type}" not found in XML configuration.')

        prefs = {}
        for child in root:
            prefs[child.tag] = child.text.strip() if child.text else ''
        return prefs

    # If there are multiple <email_report> elements
    email_configs = root.findall('email_report')

    for config in email_configs:
        report_type = config.find('report_type').text.strip()
        if report_type == email_type:
            prefs = {}
            for child in config:
                prefs[child.tag] = child.text.strip() if child.text else ''
            return prefs

    raise ValueError(f'Email type "{email_type}" not found in XML configuration.')


def capture_screen(window):
    # Determine the operating system
    os_name = platform.system()

    # Take a screenshot based on the operating system
    if os_name == 'Windows':
        # Run the Windows Snipping Tool and wait for it to finish
        subprocess.run(["SnippingTool.exe"])
        # Get the image data from the clipboard
        img = ImageGrab.grabclipboard()

        # Convert the image to bytes
        with io.BytesIO() as img_bytes:
            img.save(img_bytes, format='PNG')
            img_data = img_bytes.getvalue()

    else:
        # Unsupported operating system
        raise NotImplementedError(
            f'Operating system "{os_name}" is not supported.')

    return img_data


def save_report(report_type, patient_id, beamset_name, report_text, screenshot=None):
    """
    Save the error report document.

    Parameters:
        report_type (str): The type of report to save. values: 'error_report', 'qi_revision'
        patient_id (str): The patient ID.
        beamset_name (str): The beamset name.
        report_text (str): The text to include in the report.
        screenshot (bytes): The screenshot image data.

    Returns:
        file_path (str): The path to the saved report.
    """
    user_name = get_user_name()

    # Create a filename for the report
    now = datetime.now()
    filename = f"{patient_id}_{beamset_name}_" \
               f"{now.strftime('%Y-%m-%d_%H-%M-%S')}.docx"
    if report_type == 'error_report':
        file_path = os.path.join(ERROR_DIR, filename)
    elif report_type == 'qi_revision':
        file_path = os.path.join(QI_REPORTS_DIR, filename)

    # Embed the screenshot into the doc file
    doc = Document()

    # Set page orientation to landscape
    section = doc.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    # Set narrow margins
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    # Add report content
    doc.add_heading('Error Report')
    doc.add_paragraph(f'Patient ID: {patient_id}')
    doc.add_paragraph(f'Beamset Name: {beamset_name}')
    doc.add_paragraph(f'User Name: {user_name}')
    doc.add_paragraph(f'Time:{now.strftime("%Y-%m-%d_%H-%M-%S")}')
    doc.add_paragraph(f'Description of Issue: {report_text}')
    if screenshot:
        doc.add_picture(io.BytesIO(screenshot), width=Cm(25))

    # Save the document
    doc.save(file_path)

    # return the path
    return file_path


class EmailPackager:
    """
    Set up and send out email messages.

    Example:
    a = EmailPackager(xml_file_path='path_to_UW_Email_Preferences.xml')
    """

    def __init__(self, xml_file_path):
        self.smtp_server = None
        self.host_name = None
        self.port_num = None
        self.acct_email = None
        self.acct_password = None
        self.addr_to = None
        self.addr_from = None
        self.addr_cc = None
        self.addr_bcc = None
        self.addr_replyto = None
        self.bool_SSL = None
        self.bool_HTML = False
        self.xml_file_path = xml_file_path  # Path to the XML file

    def get_email_info(self, email_type):
        # Read from XML file
        prefs = read_email_preferences(self.xml_file_path, email_type)

        # Set attributes based on the XML values
        self.host_name = prefs.get('server')
        self.port_num = int(prefs.get('port')) if prefs.get('port') else None
        self.acct_email = prefs.get('from_email')
        self.acct_password = None  # Assuming no password in XML
        self.addr_to = list_format(prefs.get('to_email'))
        self.addr_from = self.acct_email
        self.addr_cc = list_format(prefs.get('cc_email'))
        self.addr_bcc = list_format(prefs.get('bcc_email'))
        self.addr_replyto = prefs.get('reply_to_email')
        self.bool_SSL = prefs.get('use_ssl').lower() == 'true'
        self.bool_HTML = prefs.get('use_html').lower() == 'true'

    def send_email(self,
                   email_type='error_report',
                   subj=None,
                   body=None,
                   attachments=None):
        """
        Send out email messages.

        Input:
            email_type (str): The type of email to send. Defaults to 'error_report'.
            subj (str): Subject line.
            body (str): Body text, can be plain text or HTML. Not optional.
            attachments (list): List of file paths to attach.
        """
        self.get_email_info(email_type)

        # Input validation
        if not body:
            raise ValueError('Body text message must be provided.')
        if not self.addr_to:
            raise ValueError('Recipient email address(es) must be provided.')
        if isinstance(attachments, str):
            attachments = [attachments]

        # Prepare the email message
        message = MIMEMultipart()
        message['From'] = self.addr_from
        message['To'] = ', '.join(self.addr_to)
        message['Subject'] = subj
        if self.addr_cc:
            message['Cc'] = ', '.join(self.addr_cc)
        if self.addr_replyto:
            message['Reply-To'] = self.addr_replyto

        # Attach the email body
        if self.bool_HTML:
            message.attach(MIMEText(body, 'html'))
        else:
            message.attach(MIMEText(body, 'plain'))

        # Attach files
        for f in attachments or []:
            if os.path.isfile(f):
                with open(f, "rb") as fil:
                    part = MIMEApplication(
                        fil.read(),
                        Name=os.path.basename(f)
                    )
                part['Content-Disposition'] = f'attachment; filename="{os.path.basename(f)}"'
                message.attach(part)
            else:
                print(f"Attachment {f} not found and will be skipped.")

        # Compile list of all recipients
        toaddrs = self.addr_to.copy()
        if self.addr_cc:
            toaddrs.extend(self.addr_cc)
        if self.addr_bcc:
            toaddrs.extend(self.addr_bcc)

        # Send the email
        if self.bool_SSL:
            with smtplib.SMTP_SSL(self.host_name, self.port_num) as server:
                if self.acct_password:
                    server.login(self.acct_email, self.acct_password)
                server.sendmail(self.addr_from, toaddrs, message.as_string())
        else:
            with smtplib.SMTP(self.host_name, self.port_num) as server:
                if self.acct_password:
                    server.login(self.acct_email, self.acct_password)
                server.sendmail(self.addr_from, toaddrs, message.as_string())


def test_email():
    from PlanReview.review_definitions import LOCAL_RAYSCRIPTS_DATA
    # Path to your XML configuration file
    xml_file_path = LOCAL_RAYSCRIPTS_DATA

    # Initialize the EmailPackager with the path to the XML file
    a = EmailPackager(xml_file_path=xml_file_path)

    # Send an email
    a.send_email(
        email_type='error_report',
        subj='Test Message',
        body='Hi Recipient,<br><p>You have <b>successfully</b> received this automated email.<br><p>Best regards,<br>Adam',
        attachments=None
    )


def email_report_script_error(attachment_file_path):
    from PlanReview.review_definitions import LOCAL_RAYSCRIPTS_DATA
    xml_file_path = LOCAL_RAYSCRIPTS_DATA
    email = EmailPackager(xml_file_path=xml_file_path)
    email.send_email(
        email_type='error_report',
        subj='Review Script Error Report',
        body='Please see the attached error report.',
        attachments=[attachment_file_path],
    )
    Sg.popup_ok('Error report sent to developer')


def email_report_qi_revision(attachment_file_path):
    from PlanReview.review_definitions import LOCAL_RAYSCRIPTS_DATA
    xml_file_path = LOCAL_RAYSCRIPTS_DATA
    email = EmailPackager(xml_file_path=xml_file_path)
    email.send_email(
        email_type='qi_revision',
        subj='Review Script QI Revision',
        body='Please see the attached QI revision report.',
        attachments=[attachment_file_path],
    )


if __name__ == '__main__':
    test_email()
