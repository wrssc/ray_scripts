import os
from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PlanReview.review_definitions import OUTPUT_DIR, UW_HEALTH_LOGO, LEVELS
from PlanReview.utils.get_approval_info import get_approval_info
import logging
import json

# TODO: Create header text based on input. Eliminate VMAT header thing

def set_section_columns(section, num_columns, space=Inches(0.5)):
    """
    Set the number of columns for a section
    :param section: The section to modify
    :param num_columns: The number of columns
    :param space: The space between the columns
    """
    sect_pr = section._sectPr
    columns = OxmlElement("w:cols")
    columns.set(qn("w:num"), str(num_columns))
    columns.set(qn("w:space"), str(space))
    sect_pr.append(columns)


def generate_doc(rso, tests):
    logging.info(f'tests are {json.dumps(tests)}')
    # Output file
    file_name = f"{rso.patient.PatientID}_{rso.beamset.DicomPlanLabel}.doc"
    output_file = os.path.join(OUTPUT_DIR, file_name)

    header_text = "Photon VMAT Physics Review"
    # Get approval info:
    approval_status = get_approval_info(rso.plan, rso.beamset)
    if approval_status.beamset_approved:
        current_time = str(rso.beamset.Review.ReviewTime)
    else:
        current_time = 'NA'

    demographics = {
        'Name': rso.patient.Name,
        'MRN': rso.patient.PatientID,
        'Beamset Name': rso.beamset.DicomPlanLabel,
        'Approval Time': current_time}
    # Responses to non-scriptable questions

    # Begin
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    # Add header

    header = section.header
    paragraph = header.paragraphs[0]
    # Add logo
    logo_run = paragraph.add_run()
    logo_run.add_picture(UW_HEALTH_LOGO, width=Inches(2.0), height=Inches(1.0))
    text_run = paragraph.add_run()
    text_run.text = '\t' + header_text  # For center align of text
    text_run.style = "Heading 1 Char"

    paragraph = document.add_paragraph()
    # Add Top Row Demographics
    table = document.add_table(rows=2, cols=4, style='Medium Grid 1 Accent 2')
    for index, k in enumerate(demographics):
        row_key = table.rows[0]
        row_value = table.rows[1]
        row_key.cells[index].text = k
        row_value.cells[index].text = demographics[k]
    # Add the user checks and failed
    for test_level in tests.keys():
        # Create a new section for the test level
        if document.paragraphs:
            document.add_page_break()
            document.add_section(WD_SECTION.NEW_PAGE)

        section = document.sections[-1]

        # Set the layout of the new section based on the test level
        if 'Automated' in test_level:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            set_section_columns(section, num_columns=2, space=Inches(0.5))
        else:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            set_section_columns(section, num_columns=1)

        document = add_check_list_table(tests[test_level],
                                        document, title=test_level)
    document.save(output_file)
    print('Complete')


def add_check_list_table(check_results, document, title=None):
    # Add table title if provided
    if title:
        table_title = document.add_paragraph()
        table_title.style = document.styles['Title']
        table_title.add_run(title)

    table_properties = {'NCOL': 4,
                        'WIDTH_COL': [(0, 0.25), (1, 1.), (2, 3.), (3, 3.)]}
    table = document.add_table(rows=1, cols=table_properties['NCOL'],
                               style='Light Grid Accent 2', )

    # Set row height
    row_height = Cm(0.8)

    table.add_row()
    row = table.rows[0]
    row.cells[0].text = 'Status'
    row.cells[1].text = 'Test Performed'
    row.cells[2].text = 'Result'
    row.cells[3].text = 'Reviewer Comment'
    i = 1
    for check in check_results:
        table.add_row()
        row = table.rows[i]
        icon_paragraph = row.cells[0].add_paragraph()
        icon_run = icon_paragraph.add_run()
        icon_run.add_picture(check['icon'],
                             width=Inches(0.15),
                             height=Inches(0.15))

        # Center the icon horizontally and vertically
        icon_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        icon_paragraph.space_before = Pt(0)
        icon_paragraph.space_after = Pt(0)

        row.cells[1].text = check['test_name']
        row.cells[2].text = check['result']
        row.cells[3].text = check['comment']
        i += 1

        # Set row height
        row.height = row_height

        for index, width in table_properties['WIDTH_COL']:
            for cell in table.columns[index].cells:
                cell.width = Inches(width)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Adjust font size and style for cell text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # Set font size to 10pt
                        run.font.name = 'Arial'  # Set font to Arial

    # Wrap table in a textbox for proper positioning
    textbox = document.add_textbox(Inches(0), Inches(0.5))
    textbox.width = Inches(3.5)
    textbox.height = Inches(11)
    textbox.relative_horizontal_position = 1
    textbox.vertical_anchor = 3
    textbox.relative_vertical_position = 3
    textbox.wrap_format.type = 3
    textbox.margin_left = Inches(0.25)
    textbox.margin_right = Inches(0.25)

    # Move table inside the textbox
    tbl, p = table._element, textbox._element
    tbl.getparent().remove(tbl)
    p.addnext(tbl)

    return document

def old_add_check_list_table(check_results, document, title=None):
    # Add table title if provided
    if title:
        table_title = document.add_paragraph()
        table_title.style = document.styles['Title']
        table_title.add_run(title)

    table_properties = {'NCOL': 4,
                        'WIDTH_COL': [(0, 0.25), (1, 1.), (2, 3.), (3, 3.)]}
    table = document.add_table(rows=1, cols=table_properties['NCOL'],
                               style='Light Grid Accent 2', )

    # Set row height
    row_height = Cm(0.8)

    table.add_row()
    row = table.rows[0]
    row.cells[0].text = 'Status'
    row.cells[1].text = 'Test Performed'
    row.cells[2].text = 'Result'
    row.cells[3].text = 'Reviewer Comment'
    i = 1
    for check in check_results:
        table.add_row()
        row = table.rows[i]
        icon_paragraph = row.cells[0].add_paragraph()
        icon_run = icon_paragraph.add_run()
        icon_run.add_picture(check['icon'],
                             width=Inches(0.15),
                             height=Inches(0.15))

        # Center the icon horizontally and vertically
        icon_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        icon_paragraph.space_before = Pt(0)
        icon_paragraph.space_after = Pt(0)

        row.cells[1].text = check['test_name']
        row.cells[2].text = check['result']
        row.cells[3].text = check['comment']
        i += 1

        # Set row height
        row.height = row_height

        for index, width in table_properties['WIDTH_COL']:
            for cell in table.columns[index].cells:
                cell.width = Inches(width)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Adjust font size and style for cell text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)  # Set font size to 10pt
                        run.font.name = 'Arial'  # Set font to Arial
    return document

