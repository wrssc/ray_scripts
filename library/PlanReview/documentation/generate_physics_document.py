import logging
import math
import pandas as pd
from typing import Union, Optional
from docx import Document
from docx.table import Table
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PlanReview.review_definitions import (
    LOG_DIR, UW_HEALTH_LOGO, REVIEW_LEVELS, FAIL, PASS, ALERT, RED_CIRCLE,
    GREEN_CIRCLE)
from PlanReview.utils import get_approval_info
from PlanReview.utils.constants import (
    KEY_BEAMSET, KEY_SIDE_PANEL, KEY_OUT_DOMAIN_NAME, KEY_OUT_TEST_SOURCE, SOURCE_USER, KEY_USER_COMMENT,
    KEY_OUT_TAB, KEY_OUT_RESULT, SOURCE_AUTO, KEY_OUT_DESC, KEY_OUT_MESSAGE, KEY_OUT_COMMENT, KEY_OUT_ICON,
    KEY_PROCEED_REVISE, KEY_REVISION_INFO
)
from PlanReview.utils.io_file_utils import *

# Document set up
top_margin = 0.2  # (in) top page margin
bottom_margin = 0.2  # (in) bottom page margin
left_margin = 0.25  # (in) left page margin
right_margin = 0.2  # (in) right page margin
page_height = 8.5  # (in) overall page height
row_height = 0.375  # (in) row height
domain_row_height = 0.2  # (in) domain - header row height
header_height = 0.1875  # (in) secondary header height
footer_height = 0.25  # (in) secondary footer height
table_header_height = 0.375  # (in) Table header row height
table_title_height = 0.25  # (in) Table title
table_spacing = 1.375  # (in) Table spacing
safety_margin = 0.2
# Less than this, use a page break instead of split
min_table_height = table_title_height + table_header_height + 1 * row_height


def set_section_dimensions_and_orientation(section):
    section.left_margin = Inches(left_margin)
    section.right_margin = Inches(right_margin)
    section.top_margin = Inches(top_margin)
    section.bottom_margin = Inches(bottom_margin)
    section.header_distance = Inches(top_margin)
    section.footer_distance = Inches(bottom_margin)
    # Change the orientation of the section to landscape
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def generate_doc(rso, tests, header_data, test_mode=False):
    physics_review_dir = os.path.join(LOG_DIR, "PhysicsReviews")
    patient_output_dir = os.path.join(OUTPUT_DIR, rso.patient.PatientID)
    alt_patient_output_dir = r"Q:\\RadOnc\RayStation\Reports\PhysicsReviewBetaOnly"
    patient_output_prefix = f"{rso.patient.PatientID}_" \
                            f"{rso.beamset.DicomPlanLabel}_" \
                            f"{generate_filename()}"

    if test_mode:
        latest_test_file, latest_header_file = find_latest_files(
            patient_output_dir, f"{rso.patient.PatientID}_{rso.beamset.DicomPlanLabel}_",
            ["tests.json", "header.json"])
        tests = read_tests_from_json(latest_test_file) if latest_test_file else None
        header_data = read_tests_from_json(latest_header_file) if latest_header_file else None
    else:
        test_files = [
            generate_file_path(
                patient_output_dir, patient_output_prefix, "_tests.json"),
            generate_file_path(
                physics_review_dir, patient_output_prefix, "_tests.json")
        ]
        header_files = [
            generate_file_path(
                patient_output_dir, patient_output_prefix, "_header.json"),
            generate_file_path(
                physics_review_dir, patient_output_prefix, "_header.json")
        ]
        dump_tests_to_json(tests, file_names=test_files)
        dump_tests_to_json(header_data, file_names=header_files)

    tests_df = read_data(tests)

    # Output file
    output_file = generate_file_path(
        patient_output_dir, patient_output_prefix, ".doc")
    beta_output_file = generate_file_path(
        alt_patient_output_dir, patient_output_prefix, ".doc"
    )

    # Get approval info:
    approval_status = get_approval_info(rso.plan, rso.beamset)
    if approval_status.beamset_approved:
        current_time = str(rso.beamset.Review.ReviewTime)
    else:
        current_time = 'NA'

    demographics = {
        'Name': rso.patient.Name,
        'MRN': rso.patient.PatientID,
        'Beamset Name': rso.beamset.DicomPlanLabel,
        'Approval Time': current_time}

    # Begin
    document = Document()
    section = document.sections[0]
    set_section_dimensions_and_orientation(section)
    # Add header
    add_page_header(section, rso, first_page=True)
    # Add empty footer to first page
    add_footer(section, rso, first_page=True)

    #
    # Begin body of document
    paragraph = document.add_paragraph()
    #
    # FIRST PAGE
    # Add Top Row Demographics
    table = document.add_table(rows=2, cols=4, style='Medium Grid 1 Accent 2')
    for index, k in enumerate(demographics):
        row_key = table.rows[0]
        row_value = table.rows[1]
        row_key.cells[index].text = k
        row_value.cells[index].text = demographics[k]
    # Add more front page data
    # add_simulation_data_table(document, header_data['-SIMULATION_DATA-'])
    #
    # add_treatment_instructions_table(document, header_data['-TREATMENT_INSTRUCTIONS-'])
    #
    document.add_paragraph('')  # Add spacing between sections
    add_beamset_data_table(document, header_data[KEY_BEAMSET], rso)
    add_user_comment(document, header_data[KEY_SIDE_PANEL])
    add_proceed_revise(document, header_data[KEY_SIDE_PANEL])

    #
    # USER AND AUTOMATED CHECK PAGES
    document.add_page_break()
    document.add_section(WD_SECTION.NEW_PAGE)
    # Add empty second page header
    add_page_header(section, rso, first_page=False)
    # Add footer with patient demographics
    add_footer(section, rso, first_page=False)

    available_page_height = get_usable_page()
    adjusted_page_height = available_page_height

    unique_test_levels = tests_df[KEY_OUT_TAB].unique()
    excluded_review_levels = (REVIEW_LEVELS['SANDBOX'])
    # Define custom ordering for 'RESULT'
    result_order = {FAIL: 0, ALERT: 1, PASS: 2}

    for test_level in unique_test_levels:
        if test_level in excluded_review_levels:
            continue  # Don't include these test levels in the report

        # Put in the user-entered data
        user_tl_df = tests_df[(tests_df[KEY_OUT_TEST_SOURCE] == SOURCE_USER)
                              & (tests_df[KEY_OUT_TAB] == test_level)].copy()
        # Sort DataFrame based on custom ordering
        user_tl_df['sort_key'] = user_tl_df[KEY_OUT_RESULT].map(result_order)
        user_tl_df.sort_values(by='sort_key', inplace=True)

        # Add the manual check table
        add_check_list_table(user_tl_df, document, title=f'{test_level}')
        # adjusted_page_height = add_check_table_splitting(
        #    user_tl_df, document, adjusted_page_height,f'{test_level}')
        # Auto Checks
        # Find all automated checks from source
        auto_tl_df = tests_df[(tests_df[KEY_OUT_TEST_SOURCE] == SOURCE_AUTO)
                              & (tests_df[KEY_OUT_TAB] == test_level)].copy()
        # Map the 'RESULT' values to the custom order and create a new column for sorting
        auto_tl_df['sort_key_result'] = auto_tl_df[KEY_OUT_RESULT].map(result_order)

        # Sort the DataFrame by 'KEY_OUT_DOMAIN_NAME' and the custom order
        auto_tl_df.sort_values(by=[KEY_OUT_DOMAIN_NAME, 'sort_key_result'], inplace=True)

        # Drop the temporary sorting column
        auto_tl_df.drop('sort_key_result', axis=1, inplace=True)

        # Add the automated tests
        if not auto_tl_df.empty:
            title = f'{test_level} - Automated Checks'
            add_check_list_table(auto_tl_df, document, title=f'{title}')
            # adjusted_page_height = add_check_table_splitting(
            #     auto_tl_df, document, adjusted_page_height,title)

    document.save(output_file)
    document.save(beta_output_file)
    print('Complete')


def get_usable_page():
    available_space = (page_height - top_margin - bottom_margin
                       - header_height - footer_height - safety_margin)
    return available_space


def estimate_table_length(df: pd.DataFrame, i: Optional[int] = 0, top: Optional[bool] = False) -> float:
    """
    Estimate the length of the table based on the number of rows and the
    height of each row.

    Args:
    number_of_rows (int): The number of rows in the table.

    Returns:
    float: The estimated length of the table in inches.
    """
    filtered_na_test_df = df[pd.notna(df[KEY_OUT_DOMAIN_NAME])].copy()
    unique_domain_levels = filtered_na_test_df[KEY_OUT_DOMAIN_NAME].unique().tolist()
    num_domain_rows = len(unique_domain_levels)
    number_of_rows = len(df)
    max_len_desc = df[KEY_OUT_DESC].apply(len).max()
    max_len_message = df[KEY_OUT_MESSAGE].apply(len).max()
    max_len_comment = df[KEY_OUT_COMMENT].apply(len).max()
    sum_max = max_len_comment + max_len_desc + max_len_message
    logging.debug(f'Max chars: {max_len_desc}, {max_len_message}, {max_len_comment}')
    calculated_row_height = row_height if sum_max < 240 else row_height * math.ceil(sum_max / 120)
    if top:
        space = 0
    else:
        space = table_spacing
    table_length = space + table_title_height \
                   + table_header_height \
                   + (number_of_rows - i) * calculated_row_height \
                   + num_domain_rows * domain_row_height
    return table_length


def determine_doc_type(rso):
    planning_technique = rso.beamset.PlanGenerationTechnique
    delivery_technique = rso.beamset.DeliveryTechnique
    if delivery_technique == 'DynamicArc' and planning_technique == 'Conformal':
        header_text = 'Photon Conformal Arc'
    elif delivery_technique == 'DynamicArc' and planning_technique == 'Imrt':
        header_text = 'Photon VMAT'
    elif delivery_technique == 'ApplicatorAndCutout':
        header_text = 'Electron'
    elif delivery_technique == 'TomoHelical':
        if 'T3D' not in rso.beamset.DicomPlanLabel:
            header_text = 'TomoHelical'
        else:
            header_text = 'Tomo3D'
    elif delivery_technique == 'SMLC':
        if planning_technique == 'Imrt':
            header_text = 'Static Field IMRT'
        else:
            header_text = '3D'
    else:
        header_text = 'Unsupported Technique'
    header_text = header_text + ' Physics Review'
    return header_text


def add_page_header(section, rso, first_page=False):
    # Add header
    header = section.header
    header_paragraph = header.paragraphs[0]

    if first_page:
        # Add logo
        logo_width = 2.5
        logo_height = logo_width * 240 / 920
        logo_run = header_paragraph.add_run()
        _ = logo_run.add_picture(UW_HEALTH_LOGO,
                                 width=Inches(logo_width),
                                 height=Inches(logo_height))
        # Get document header title
        header_text = '\t' * 2 + determine_doc_type(rso)
    else:
        header.is_linked_to_previous = False
        header_text = ""
    header_run = header_paragraph.add_run()
    header_run.text = header_text
    header_run.style = "Heading 1 Char"


def add_footer(section, rso, first_page=False):
    demographics = {
        'Name': rso.patient.Name,
        'MRN': rso.patient.PatientID,
        'Beamset Name': rso.beamset.DicomPlanLabel, }
    if first_page:
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_run = footer_paragraph.add_run()
        footer_run.text = ""  # no footer content for first page
    else:
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_paragraph = footer.paragraphs[0]
        footer_run = footer_paragraph.add_run()
        footer_run.text = "\t".join(
            [key + ':' + value for key, value in demographics.items()])
        footer_run.style = "Heading 4 Char"


def add_check_list_table(df: pd.DataFrame, document: Document,
                         title: str = None) -> Table:
    """
    Adds a table to the document for the given check results. Page breaks are
    added only if the table does not fit on the current page.

    Args:
    check_results (list): List of check results to be added to the table.
    document (Document): The document to add the table to.
    title (str, optional): The title of the table. Defaults to None.

    Returns:
    Table: The created table.
    :type df: pd.Dataframe
    """
    # Set row height
    table_row_height = Inches(row_height)
    # Calculate estimated table length
    table_length = estimate_table_length(df)

    if title:
        table_title = document.add_paragraph()
        table_title.style = document.styles['Heading 1']
        table_title.add_run(title)

    # Calculate the available width for the table
    page_width = document.sections[-1].page_width.inches
    available_width = page_width - left_margin - right_margin - 0.5  # Subtract 0.5 inches for

    # Set the column widths (widths are in proportion to the amount of text)
    long_keys = [KEY_OUT_DESC, KEY_OUT_MESSAGE, KEY_OUT_COMMENT]
    # Determine the maximum length of text for each key
    col_text_width = [df[col].str.len().max() for col in long_keys]
    total_text_width = sum(col_text_width)

    # Calculate proportional column widths
    col_widths = [0.5] + [
        (text_width / total_text_width) * available_width
        for text_width in col_text_width
    ]

    table_properties = {'NCOL': 4, 'WIDTH_COL': [
        (i, Inches(w)) for i, w in enumerate(col_widths)
    ]}

    table = document.add_table(
        rows=1, cols=table_properties['NCOL'], style='Light Grid Accent 2')
    # Set the header attribute to True to repeat column headers
    table.header = True

    row = table.rows[0]
    row.height = table_row_height
    row.cells[0].text = 'Status'
    row.cells[1].text = 'Test Performed'
    row.cells[2].text = 'Result'
    row.cells[3].text = 'Reviewer Comment'
    i = 1

    last_domain = None
    for _, check in df.iterrows():
        current_domain = check[KEY_OUT_DOMAIN_NAME]
        # Skip NaN domains and check for new domain names
        if pd.notna(current_domain) and current_domain != last_domain:
            add_domain_row(table, current_domain, i)
            last_domain = current_domain
            i += 1

        table.add_row()
        row = table.rows[i]
        icon_paragraph = row.cells[0].paragraphs[0]
        icon_run = icon_paragraph.add_run()
        icon_run.add_picture(
            check[KEY_OUT_ICON], width=Inches(0.15), height=Inches(0.15))
        icon_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].text = check[KEY_OUT_DESC]
        row.cells[2].text = check[KEY_OUT_MESSAGE]
        row.cells[3].text = check[KEY_OUT_COMMENT]
        i += 1

        # Set row height
        row.height = table_row_height

        for index, width in table_properties['WIDTH_COL']:
            if index < len(table.columns):  # Ensure the index is within the valid range
                for cell in table.columns[index].cells:
                    cell.width = width
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # Adjust font size and style for cell text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)  # Set font size to 10pt
                            run.font.name = 'Arial'  # Set font to Arial

    return table


def add_table_row(table: Table, data: Union[list, tuple]) -> None:
    """
    Adds a row to the table with specified data.

    Args:
        table (Table): The table to which the row will be added.
        data (Union[list, tuple]): The data to populate the row.
    """
    row = table.add_row()
    for i, value in enumerate(data):
        row.cells[i].text = str(value)
        row.cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_domain_row(table, domain, index):
    """
    Adds a domain row to the table with a specific background color.

    Args:
        table (Table): The table to which the domain row will be added.
        domain (str): The domain name to display in the row.
    """
    table.add_row()
    row = table.rows[index]
    row.cells[0].merge(row.cells[-1])
    row.cells[0].text = domain
    row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Set background color (shade) to light gray; 'auto' could be replaced with a specific color
    add_shading_to_cell(row.cells[0], 'D3D3D3')


def add_simulation_data_table(doc, data):
    simulation_table = doc.add_table(rows=1, cols=2)
    simulation_table.style = 'Table Grid'
    header_cells = simulation_table.rows[0].cells
    header_cells[0].text = 'Key'
    header_cells[1].text = 'Value'

    for key, value in data.items():
        add_table_row(simulation_table, (key, value))

    doc.add_paragraph('')  # Add spacing between sections


def add_user_comment(doc, data):
    comment_table = doc.add_table(rows=2, cols=1)
    comment_table.style = 'Light List Accent 2'
    comment_table.cell(0, 0).text = 'Physicist Comments'
    comment_table.cell(1, 0).text = str(data[KEY_USER_COMMENT])
    doc.add_paragraph('')  # Add spacing between sections


def row_proceed_revise(data):
    status_mapping = {
        "Revise": ("Revise", str(data.get(KEY_REVISION_INFO, "")), RED_CIRCLE),
        "Proceed": ("Proceed","", GREEN_CIRCLE),
        "QIProceed": ("Proceed","", GREEN_CIRCLE),
    }

    recommendation, comment, icon = status_mapping.get(data[KEY_PROCEED_REVISE], ("", "", ""))

    if not comment and not icon:
        logging.warning(f'Unknown plan revision outcome {data[KEY_PROCEED_REVISE]}')

    return recommendation, comment, icon


def add_proceed_revise(doc, data):
    # Build the revision status table
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Light List Accent 2'
    table.rows[1].height = Inches(row_height)
    # Set individual column widths
    page_width = doc.sections[-1].page_width.inches
    w_col1 = Inches(0.5)
    w_col2 = Inches(1.5)
    w_col3 = Inches(page_width) - w_col2 - w_col1
    col_widths = [w_col1, w_col2, w_col3]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_widths[i]

    # Fill in table headers
    table.cell(0, 0).text = 'Status'
    table.cell(0, 1).text = 'Recommendation'
    table.cell(0, 2).text = 'Status Comment'

    # Check status
    recommendation, comment, icon = row_proceed_revise(data)

    # Add icon
    icon_cell = table.cell(1,0)
    icon_paragraph = icon_cell.paragraphs[0]
    icon_run = icon_paragraph.add_run()
    icon_run.add_picture(icon, width=Inches(0.20), height=Inches(0.20))
    icon_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    icon_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add Recommendation and make it bold
    rec_cell = table.cell(1, 1)
    rec_paragraph = rec_cell.paragraphs[0]
    rec_run = rec_paragraph.add_run(recommendation)
    rec_run.bold = True
    rec_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add Comments
    table.cell(1, 2).text = comment
    table.cell(1, 2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add spacing between sections
    doc.add_paragraph('')


def add_treatment_instructions_table(doc, data):
    instructions_table = doc.add_table(rows=1, cols=3)
    instructions_table.style = 'Table Grid'
    header_cells = instructions_table.rows[0].cells
    header_cells[0].text = 'Radio Button State'
    header_cells[1].text = 'Selected Option'
    header_cells[2].text = 'Value'

    for key, value in data.items():
        if key[0].startswith('-INSTRUCTION--RADIO-'):
            radio_state = key[0].split('-')[-1]
            index = key[1]
            combo_key = ('-INSTRUCTION--COMBO-', index)
            input_text_key = ('-INSTRUCTION--INPUT-TEXT-', index)
            if combo_key in data:
                selected_option = data[combo_key]
            elif input_text_key in data:
                selected_option = data[input_text_key]
            else:
                selected_option = ''
            add_table_row(instructions_table,
                          (radio_state, selected_option, value))

    doc.add_paragraph('')  # Add spacing between sections


def add_beamset_data_table(doc, data, rso):
    logging.debug(f'Beamset keys {data.keys()}')
    beamset_count = data["-BEAMSET--COUNT-"]
    beamset_table = doc.add_table(rows=1, cols=2)
    beamset_table.style = 'Medium Grid 1 Accent 2'

    # Set the first column width to 1.5 inches
    for cell in beamset_table.column_cells(0):
        cell.width = Inches(1.5)

    # Set the headers for the main table
    header_cells = beamset_table.rows[0].cells
    header_cells[0].text = 'Beamset'
    header_cells[1].text = 'Beamset Details'

    # Loop through beamsets
    for i in range(beamset_count):
        beamset_number = ("-BEAMSET_SELECT-", i)
        beamset_name = data[beamset_number]
        row = beamset_table.add_row()
        row.cells[0].text = beamset_name

        # Add a nested table for beamset details
        nested_table = row.cells[1].add_table(rows=3, cols=2)
        nested_table.style = 'Medium Grid 2 Accent 2'

        # Add Number of Fractions row
        nested_table.cell(0, 0).text = "Number of Fractions"
        nested_table.cell(0, 1).text = str(data[("-BEAMSET--N_FRACTIONS-", i)])

        # Add Approval Status row
        approval_status = get_approval_info(
            rso.plan, rso.plan.BeamSets[beamset_name])
        if approval_status.beamset_approved:
            approval_date = str(
                rso.plan.BeamSets[beamset_name].Review.ReviewTime)
        else:
            approval_date = 'NA'
        nested_table.cell(1, 0).text = "Beamset Approval Date"
        nested_table.cell(1, 1).text = approval_date

        # Add Targets nested table
        targets_table = nested_table.cell(2, 1).add_table(rows=1, cols=3)
        targets_table.style = 'Medium Grid 2 Accent 2'

        # Set the headers for the targets table
        header_cells = targets_table.rows[0].cells
        header_cells[0].text = 'Target Name'
        header_cells[1].text = 'Dose per Fraction (Gy)'
        header_cells[2].text = 'Total Dose (Gy)'

        # Loop through targets
        target_count = data[("-BEAMSET--TARGET_COUNT-", i)]
        for j in range(target_count):
            target_name = data[("-BEAMSET--TARGET-NAME", i, j)]
            fraction_dose = data[("-BEAMSET--FRACTION-DOSE-", i, j)]
            total_dose = data[("-BEAMSET--DOSE-", i, j)]
            add_table_row(targets_table, (target_name,
                                          fraction_dose, total_dose))

    # Add spacing between sections
    doc.add_paragraph('')





#
# TODO: Optional export to pdf
# from docx2pdf import convert
# file_name = f"{rso.patient.PatientID}_{rso.beamset.DicomPlanLabel}.docx"  # Change the file
# output_file = os.path.join(OUTPUT_DIR, file_name)
# pdf_file_name = f"{rso.patient.PatientID}_{rso.beamset.DicomPlanLabel}.pdf"
# pdf_output_file = os.path.join(OUTPUT_DIR, pdf_file_name)
# document.save(output_file)  # Save the document as a DOCX file
# convert(output_file, pdf_output_file)  # Convert the DOCX file to a PDF file
# os.remove(output_file)  # Remove the DOCX file after conversion, if needed





def read_data(data):
    return pd.DataFrame(data)


def add_shading_to_cell(cell, color):
    # Create new shading element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)

    # Add shading to cell properties
    cell._tc.get_or_add_tcPr().append(shd)


